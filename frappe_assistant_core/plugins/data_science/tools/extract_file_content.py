# -*- coding: utf-8 -*-
# Frappe Assistant Core - AI Assistant integration for Frappe Framework
# Copyright (C) 2025 Paul Clinton
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

"""
File Content Extraction Tool for Data Science Plugin.
Extracts content from various file formats (PDF, images, CSV, Excel, documents) for LLM processing.
"""

import frappe
from frappe import _
import io
import base64
import mimetypes
from typing import Dict, Any, Optional
from frappe_assistant_core.core.base_tool import BaseTool
import os


class ExtractFileContent(BaseTool):
    """
    📄 File Content Extraction Tool for LLM Processing
    
    Extract content from various file formats and prepare it for LLM analysis 
    through MCP tools.
    
    📁 **SUPPORTED FORMATS**:
    • PDFs - Text extraction, OCR for scanned docs, table extraction
    • Images - OCR text extraction, object detection, content analysis
    • CSV/Excel - Data parsing, validation, transformation, insights
    • Documents - DOCX, TXT content extraction and analysis
    
    🎯 **KEY CAPABILITIES**:
    • Text extraction from any document type
    • OCR for scanned documents and images
    • Table and structured data extraction
    • Format-aware content parsing
    • Data validation for CSV/Excel
    • Multi-language support for OCR
    • Content preparation for LLM processing
    
    💡 **USE CASES**:
    • Extract invoice content for LLM analysis
    • Read contracts and legal documents
    • Extract data from scanned forms
    • Parse CSV/Excel data for processing
    • OCR scanned documents
    • Prepare documents for Q&A with LLMs
    • Extract structured data for validation
    """
    
    def __init__(self):
        super().__init__()
        self.name = "extract_file_content"
        self.description = self._get_description()
        self.requires_permission = "File"  # Requires permission to access File DocType
        
        self.inputSchema = {
            "type": "object",
            "properties": {
                "file_url": {
                    "type": "string",
                    "description": "File URL from Frappe (e.g., '/files/invoice.pdf' or '/private/files/document.docx'). Provide either file_url OR file_name."
                },
                "file_name": {
                    "type": "string",
                    "description": "Alternative: File name from File DocType (e.g., 'invoice-2024.pdf'). Provide either file_url OR file_name."
                },
                "operation": {
                    "type": "string",
                    "enum": ["extract", "ocr", "parse_data", "extract_tables"],
                    "description": "Operation: 'extract' (get text/data), 'ocr' (extract text from images), 'parse_data' (structured data from CSV/Excel), 'extract_tables' (extract tables from PDFs)"
                },
                "language": {
                    "type": "string",
                    "default": "eng",
                    "description": "OCR language code (eng, fra, deu, spa, etc.)"
                },
                "output_format": {
                    "type": "string",
                    "enum": ["json", "text", "markdown"],
                    "default": "text",
                    "description": "Output format for extracted content"
                },
                "max_pages": {
                    "type": "integer",
                    "default": 50,
                    "description": "Maximum pages to process for PDFs"
                }
            },
            "required": ["operation"]
        }
    
    def _get_description(self) -> str:
        """Get tool description"""
        return """📄 **Extract File Content** - Extract text and data from various file formats for LLM processing.

📁 **Supported Formats**: PDF, JPG/PNG (OCR), CSV/Excel, DOCX/TXT
🔍 **Operations**: Extract text, OCR images, parse structured data, extract tables
🎯 **Use Cases**: Read invoices, contracts, forms, spreadsheets for LLM analysis

💡 **Examples**:
• Extract text from PDF: operation='extract', file_url='/files/contract.pdf'
• OCR scan: operation='ocr', file_url='/files/invoice_scan.jpg'
• Parse CSV data: operation='parse_data', file_url='/files/data.csv'
• Extract PDF tables: operation='extract_tables', file_url='/files/report.pdf'"""
    
    def execute(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Execute file content extraction"""
        try:
            # Validate dependencies first
            dep_check = self._check_dependencies()
            if not dep_check["success"]:
                return dep_check
            
            # Get file from Frappe
            file_doc = self._get_file_document(arguments)
            if not file_doc:
                return {
                    "success": False,
                    "error": "File not found or access denied"
                }
            
            # Check file size limits
            if not self._check_file_size(file_doc):
                return {
                    "success": False,
                    "error": f"File size exceeds limit of 50MB"
                }
            
            # Get file content
            file_content = self._get_file_content(file_doc)
            if not file_content:
                return {
                    "success": False,
                    "error": "Failed to read file content"
                }
            
            # Detect file type
            file_type = self._detect_file_type(file_doc)
            
            # Process based on operation
            operation = arguments.get("operation", "extract")
            
            if operation == "extract":
                result = self._extract_content(file_content, file_type, arguments)
            elif operation == "ocr":
                result = self._perform_ocr(file_content, arguments)
            elif operation == "parse_data":
                if file_type in ['csv', 'excel']:
                    result = self._extract_content(file_content, file_type, arguments)
                else:
                    return {
                        "success": False,
                        "error": "parse_data operation only supports CSV and Excel files"
                    }
            elif operation == "extract_tables":
                if file_type == 'pdf':
                    result = self._extract_pdf_tables(file_content, arguments)
                else:
                    return {
                        "success": False,
                        "error": "extract_tables operation only supports PDF files"
                    }
            else:
                return {
                    "success": False,
                    "error": f"Unknown operation: {operation}"
                }
            
            # Add file metadata to result
            if result.get("success"):
                result["file_info"] = {
                    "name": file_doc.file_name,
                    "type": file_type,
                    "size": file_doc.file_size if hasattr(file_doc, 'file_size') else len(file_content),
                    "url": file_doc.file_url
                }
            
            return result
            
        except Exception as e:
            frappe.log_error(
                title="File Processing Error",
                message=f"Error processing file: {str(e)}"
            )
            return {
                "success": False,
                "error": str(e)
            }
    
    def _check_dependencies(self) -> Dict[str, Any]:
        """Check if required dependencies are available"""
        missing_deps = []
        
        # Check required libraries
        try:
            import PyPDF2
        except ImportError:
            missing_deps.append("PyPDF2")
        
        try:
            import PIL
        except ImportError:
            missing_deps.append("Pillow")
        
        try:
            import pandas
        except ImportError:
            missing_deps.append("pandas")
        
        if missing_deps:
            return {
                "success": False,
                "error": f"Missing dependencies: {', '.join(missing_deps)}. Please install them using pip."
            }
        
        return {"success": True}
    
    def _get_file_document(self, arguments: Dict[str, Any]) -> Optional[Any]:
        """Get file document from Frappe"""
        file_url = arguments.get("file_url")
        file_name = arguments.get("file_name")
        
        try:
            if file_url:
                # Get file by URL
                file_doc = frappe.get_all(
                    "File",
                    filters={"file_url": file_url},
                    fields=["*"],
                    limit=1
                )
                if file_doc:
                    return frappe.get_doc("File", file_doc[0].name)
            
            elif file_name:
                # Get file by name
                file_doc = frappe.get_all(
                    "File",
                    filters={"file_name": file_name},
                    fields=["*"],
                    limit=1
                )
                if file_doc:
                    return frappe.get_doc("File", file_doc[0].name)
            
            return None
            
        except Exception as e:
            frappe.log_error(f"Error getting file document: {str(e)}")
            return None
    
    def _check_file_size(self, file_doc) -> bool:
        """Check if file size is within limits"""
        max_size = 50 * 1024 * 1024  # 50MB
        
        try:
            if hasattr(file_doc, 'file_size') and file_doc.file_size:
                return file_doc.file_size <= max_size
            
            # Try to get file size from file system
            file_path = self._get_file_path(file_doc)
            if file_path and os.path.exists(file_path):
                return os.path.getsize(file_path) <= max_size
            
            return True  # Allow if we can't determine size
            
        except Exception:
            return True
    
    def _get_file_path(self, file_doc) -> Optional[str]:
        """Get absolute file path"""
        if file_doc.file_url:
            if file_doc.file_url.startswith("/private"):
                # Private file
                return frappe.get_site_path(file_doc.file_url.lstrip("/"))
            elif file_doc.file_url.startswith("/files"):
                # Public file
                return frappe.get_site_path("public", file_doc.file_url.lstrip("/"))
        return None
    
    def _get_file_content(self, file_doc) -> Optional[bytes]:
        """Get file content as bytes"""
        try:
            file_path = self._get_file_path(file_doc)
            if file_path and os.path.exists(file_path):
                with open(file_path, 'rb') as f:
                    return f.read()
            
            # Try to get from file_doc content if stored
            if hasattr(file_doc, 'content'):
                if isinstance(file_doc.content, str):
                    # Base64 encoded content
                    return base64.b64decode(file_doc.content)
                return file_doc.content
            
            return None
            
        except Exception as e:
            frappe.log_error(f"Error reading file content: {str(e)}")
            return None
    
    def _detect_file_type(self, file_doc) -> str:
        """Detect file type from document"""
        file_name = file_doc.file_name or file_doc.file_url or ""
        file_name_lower = file_name.lower()
        
        if file_name_lower.endswith('.pdf'):
            return 'pdf'
        elif file_name_lower.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
            return 'image'
        elif file_name_lower.endswith(('.csv', '.tsv')):
            return 'csv'
        elif file_name_lower.endswith(('.xlsx', '.xls')):
            return 'excel'
        elif file_name_lower.endswith('.docx'):
            return 'docx'
        elif file_name_lower.endswith(('.txt', '.text')):
            return 'text'
        else:
            # Try to detect from MIME type
            mime_type, _ = mimetypes.guess_type(file_name)
            if mime_type:
                if 'pdf' in mime_type:
                    return 'pdf'
                elif 'image' in mime_type:
                    return 'image'
                elif 'csv' in mime_type or 'tab-separated' in mime_type:
                    return 'csv'
                elif 'spreadsheet' in mime_type or 'excel' in mime_type:
                    return 'excel'
                elif 'word' in mime_type:
                    return 'docx'
                elif 'text' in mime_type:
                    return 'text'
            
            return 'unknown'
    
    def _extract_content(self, file_content: bytes, file_type: str, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Extract content based on file type"""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_content(file_content, arguments)
            elif file_type == 'image':
                return self._extract_image_content(file_content, arguments)
            elif file_type == 'csv':
                return self._extract_csv_content(file_content)
            elif file_type == 'excel':
                return self._extract_excel_content(file_content)
            elif file_type == 'docx':
                return self._extract_docx_content(file_content)
            elif file_type == 'text':
                return self._extract_text_content(file_content)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_type}"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"Content extraction failed: {str(e)}"
            }
    
    def _extract_pdf_content(self, file_content: bytes, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Extract content from PDF"""
        try:
            import PyPDF2
            from PyPDF2 import PdfReader
            
            # Create PDF reader
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            max_pages = arguments.get("max_pages", 50)
            num_pages = min(len(reader.pages), max_pages)
            
            # Extract text from each page
            text_content = []
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            combined_text = "\n\n".join(text_content)
            
            # If no text extracted, might be scanned PDF
            if not combined_text.strip():
                return {
                    "success": True,
                    "content": "",
                    "message": "No text found in PDF. This might be a scanned document. Use operation='ocr' for scanned PDFs.",
                    "pages": num_pages
                }
            
            return {
                "success": True,
                "content": combined_text,
                "pages": num_pages,
                "extracted_pages": len(text_content)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF extraction error: {str(e)}"
            }
    
    def _extract_image_content(self, file_content: bytes, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Extract content from image using OCR"""
        return self._perform_ocr(file_content, arguments)
    
    def _perform_ocr(self, file_content: bytes, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Perform OCR on image content"""
        try:
            # Check if pytesseract is available
            try:
                import pytesseract
                from PIL import Image
            except ImportError:
                return {
                    "success": False,
                    "error": "OCR dependencies not installed. Please install pytesseract and Pillow.",
                    "install_command": "pip install pytesseract pillow"
                }
            
            # Check if tesseract is installed on system
            try:
                pytesseract.get_tesseract_version()
            except Exception:
                return {
                    "success": False,
                    "error": "Tesseract OCR not installed on system. Please install tesseract-ocr.",
                    "install_command": "sudo apt-get install tesseract-ocr (Linux) or brew install tesseract (Mac)"
                }
            
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            # Perform OCR
            language = arguments.get("language", "eng")
            extracted_text = pytesseract.image_to_string(image, lang=language)
            
            if not extracted_text.strip():
                return {
                    "success": True,
                    "content": "",
                    "message": "No text detected in image"
                }
            
            return {
                "success": True,
                "content": extracted_text,
                "ocr_language": language
            }
            
        except Exception as e:
            # Fallback message if OCR fails
            return {
                "success": True,
                "content": "[OCR not available - image file detected]",
                "message": f"OCR failed: {str(e)}. To enable OCR, install tesseract-ocr system package.",
                "fallback": True
            }
    
    def _extract_csv_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from CSV"""
        try:
            import pandas as pd
            
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    df = pd.read_csv(io.BytesIO(file_content), encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return {
                    "success": False,
                    "error": "Failed to decode CSV file with common encodings"
                }
            
            # Convert to dict for serialization
            data_dict = {
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "sample_data": df.head(10).to_dict('records'),
                "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
            
            # Create text representation
            text_content = f"CSV Data Summary:\n"
            text_content += f"Columns: {', '.join(data_dict['columns'])}\n"
            text_content += f"Total Rows: {data_dict['row_count']}\n\n"
            text_content += "Sample Data:\n"
            text_content += df.head(10).to_string()
            
            return {
                "success": True,
                "content": text_content,
                "structured_data": data_dict
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"CSV extraction error: {str(e)}"
            }
    
    def _extract_excel_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from Excel"""
        try:
            import pandas as pd
            
            # Read Excel file
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
            
            all_sheets_content = []
            structured_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Store structured data
                structured_data[sheet_name] = {
                    "columns": df.columns.tolist(),
                    "row_count": len(df),
                    "sample_data": df.head(10).to_dict('records')
                }
                
                # Create text representation
                sheet_content = f"=== Sheet: {sheet_name} ===\n"
                sheet_content += f"Columns: {', '.join(df.columns.tolist())}\n"
                sheet_content += f"Rows: {len(df)}\n\n"
                sheet_content += df.head(10).to_string()
                
                all_sheets_content.append(sheet_content)
            
            combined_content = "\n\n".join(all_sheets_content)
            
            return {
                "success": True,
                "content": combined_content,
                "structured_data": structured_data,
                "sheet_count": len(excel_file.sheet_names)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Excel extraction error: {str(e)}"
            }
    
    def _extract_docx_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from DOCX"""
        try:
            # Check if python-docx is available
            try:
                from docx import Document
            except ImportError:
                return {
                    "success": False,
                    "error": "python-docx not installed. Please install it using: pip install python-docx"
                }
            
            # Read document
            doc = Document(io.BytesIO(file_content))
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables if any
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine content
            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n=== Tables ===\n\n" + "\n\n".join(tables_text)
            
            return {
                "success": True,
                "content": content,
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX extraction error: {str(e)}"
            }
    
    def _extract_text_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from text file"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
                try:
                    text = file_content.decode(encoding)
                    return {
                        "success": True,
                        "content": text,
                        "encoding": encoding
                    }
                except UnicodeDecodeError:
                    continue
            
            return {
                "success": False,
                "error": "Failed to decode text file with common encodings"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Text extraction error: {str(e)}"
            }
    
    def _extract_pdf_tables(self, file_content: bytes, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Extract tables from PDF"""
        try:
            # Try using pdfplumber for better table extraction
            try:
                import pdfplumber
                import pandas as pd
                
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    all_tables = []
                    max_pages = min(arguments.get("max_pages", 50), len(pdf.pages))
                    
                    for page_num in range(max_pages):
                        page = pdf.pages[page_num]
                        tables = page.extract_tables()
                        
                        for table_idx, table in enumerate(tables):
                            if table:
                                # Convert to DataFrame for better structure
                                df = pd.DataFrame(table[1:], columns=table[0] if table else None)
                                all_tables.append({
                                    "page": page_num + 1,
                                    "table_index": table_idx + 1,
                                    "data": df.to_dict('records'),
                                    "rows": len(df),
                                    "columns": len(df.columns)
                                })
                    
                    if not all_tables:
                        return {
                            "success": True,
                            "message": "No tables found in PDF",
                            "tables": []
                        }
                    
                    return {
                        "success": True,
                        "tables": all_tables,
                        "total_tables": len(all_tables),
                        "pages_processed": max_pages
                    }
                    
            except ImportError:
                # Fallback to basic extraction if pdfplumber not available
                return {
                    "success": True,
                    "message": "Table extraction requires pdfplumber. Install with: pip install pdfplumber",
                    "fallback": True
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"Table extraction error: {str(e)}"
            }


# Make sure class is available for discovery
# The plugin manager will find ExtractFileContent automatically